"""
P7b — pull plain text out of an uploaded document for the internal KB.

    extract("policy.pdf", data)  -> ("policy", "<text>")

.txt / .md / .csv / .json decode directly; .pdf needs `pypdf`, .docx needs
`python-docx` (both optional — a clear error if the file type needs one that
isn't installed). Text is capped so one huge upload can't blow the embedder.
"""

from __future__ import annotations

import io
import pathlib

MAX_BYTES = 8 * 1024 * 1024        # 8 MB upload
MAX_CHARS = 200_000               # ~50 KB of chunks after embedding

_TEXT_EXT = {".txt", ".md", ".markdown", ".csv", ".json", ".rst", ".log"}


def extract(filename: str, data: bytes) -> tuple[str, str]:
    if len(data) > MAX_BYTES:
        raise ValueError(f"file is {len(data) // 1024} KB — the limit is {MAX_BYTES // 1024} KB")
    ext = pathlib.Path(filename or "").suffix.lower()
    title = (pathlib.Path(filename or "upload").stem or "upload")[:200]

    if ext in _TEXT_EXT or not ext:
        text = data.decode("utf-8", errors="replace")
    elif ext == ".pdf":
        text = _pdf(data)
    elif ext == ".docx":
        text = _docx(data)
    else:
        raise ValueError(f"unsupported file type {ext or '(none)'} — "
                         "use pdf, docx, md, txt, csv or json")

    text = text.replace("\x00", "").strip()
    if not text:
        raise ValueError("no text could be extracted from the file")
    return title, text[:MAX_CHARS]


def _pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("PDF upload needs `pip install pypdf`") from e
    reader = PdfReader(io.BytesIO(data))
    return "\n\n".join((p.extract_text() or "") for p in reader.pages)


def _docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("DOCX upload needs `pip install python-docx`") from e
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs]
    for tbl in d.tables:
        for row in tbl.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)
